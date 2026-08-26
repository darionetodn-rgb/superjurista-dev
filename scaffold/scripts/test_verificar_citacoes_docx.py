#!/usr/bin/env python3
"""Teste dirigido do suporte a .docx do verificar_citacoes.py (13/08/2026).

Origem: linha do backlog da Auditoria CC (F3) — "verificar_citacoes.py abre
.docx como texto (UnicodeDecodeError)"; a minuta do escritorio e sempre .docx e
a conversao manual se repetia a cada corrida do /pipeline-revisao-minuta.

Os 5 casos cobrem o que o conserto tem de garantir, e cada um FALHA se a
respectiva perna quebrar:

  A  .docx com citacao que consta do corpus              -> exit 0
  B  .docx com citacao que NAO consta                    -> exit 1
  C  citacao dentro de TABELA que nao consta             -> exit 1
     (documento.paragraphs nao enxerga tabela: sem o percurso pelo XML este
      caso passaria como se estivesse limpo — falso verde)
  D  .md segue funcionando                               -> nao-regressao
  E  --ignorar-apos corta em Heading do .docx            -> exit 0

Roda sozinho, sem framework: py -3 scripts/test_verificar_citacoes_docx.py
"""

import os
import shutil
import subprocess
import sys
import tempfile

AQUI = os.path.dirname(os.path.abspath(__file__))
SCRIPT = os.path.join(AQUI, "verificar_citacoes.py")

# >= LIMIAR (60) chars normalizados, senao fica fora do regime do gate
TRECHO_BOM = ("a prescricao intercorrente em execucao fiscal exige a previa "
              "intimacao da Fazenda Publica sobre o arquivamento dos autos")
TRECHO_MAU = ("o contribuinte esta dispensado de qualquer garantia do juizo "
              "quando a certidao de divida ativa contiver vicio meramente formal")


def _monta_workspace(tmp, texto_dos_autos):
    ws = os.path.join(tmp, "0000000-00.0000.0.00.0000")
    os.makedirs(ws, exist_ok=True)
    with open(os.path.join(ws, "processo.txt"), "w", encoding="utf-8") as f:
        f.write(texto_dos_autos)
    return ws


def _docx(caminho, paragrafos=(), linhas_tabela=(), headings=()):
    import docx
    documento = docx.Document()
    for texto, estilo in headings:
        documento.add_paragraph(texto, style=estilo)
    for p in paragrafos:
        documento.add_paragraph(p)
    if linhas_tabela:
        tabela = documento.add_table(rows=len(linhas_tabela), cols=1)
        for i, conteudo in enumerate(linhas_tabela):
            tabela.rows[i].cells[0].text = conteudo
    documento.save(caminho)
    return caminho


def _roda(ws, doc, *extra):
    r = subprocess.run([sys.executable, SCRIPT, ws, "--doc", doc, *extra],
                       capture_output=True, text=True)
    return r.returncode, (r.stdout or "") + (r.stderr or "")


def main():
    tmp = tempfile.mkdtemp(prefix="test_citacoes_docx_")
    falhas = []
    try:
        ws = _monta_workspace(tmp, TRECHO_BOM)

        # A — .docx cuja citacao consta dos autos
        doc = _docx(os.path.join(tmp, "a.docx"),
                    paragrafos=[f'Conforme os autos, "{TRECHO_BOM}".'])
        codigo, saida = _roda(ws, doc)
        if codigo != 0:
            falhas.append(f"A: esperado exit 0, veio {codigo}\n{saida}")

        # B — .docx com citacao sem lastro
        doc = _docx(os.path.join(tmp, "b.docx"),
                    paragrafos=[f'Alega-se que "{TRECHO_MAU}".'])
        codigo, saida = _roda(ws, doc)
        if codigo != 1:
            falhas.append(f"B: esperado exit 1, veio {codigo}\n{saida}")

        # C — citacao sem lastro DENTRO DE TABELA (a perna que paragraphs perde)
        doc = _docx(os.path.join(tmp, "c.docx"),
                    paragrafos=["Quadro-resumo das teses:"],
                    linhas_tabela=[f'Tese 1: "{TRECHO_MAU}"'])
        codigo, saida = _roda(ws, doc)
        if codigo != 1:
            falhas.append("C: citacao em TABELA passou despercebida (falso verde) — "
                          f"esperado exit 1, veio {codigo}\n{saida}")

        # D — nao-regressao: .md continua valendo
        caminho_md = os.path.join(tmp, "d.md")
        with open(caminho_md, "w", encoding="utf-8") as f:
            f.write(f'Texto da minuta: "{TRECHO_MAU}".')
        codigo, saida = _roda(ws, caminho_md)
        if codigo != 1:
            falhas.append(f"D: .md regrediu — esperado exit 1, veio {codigo}\n{saida}")

        # E — --ignorar-apos corta em Heading do .docx (a citacao ruim fica fora)
        doc = _docx(os.path.join(tmp, "e.docx"),
                    headings=[("Log de Alteracoes", "Heading 1")],
                    paragrafos=[f'Auto-citacao do log: "{TRECHO_MAU}".'])
        codigo, saida = _roda(ws, doc, "--ignorar-apos", "log de alteracoes")
        if codigo != 0:
            falhas.append("E: --ignorar-apos nao cortou no Heading do .docx — "
                          f"esperado exit 0, veio {codigo}\n{saida}")
    finally:
        shutil.rmtree(tmp, ignore_errors=True)

    if falhas:
        print("FALHOU:")
        for f in falhas:
            print("  - " + f.replace("\n", "\n    "))
        return 1
    print("TODOS OS 5 CASOS PASSARAM (A .docx ok, B sem lastro, C tabela, "
          "D .md nao-regressao, E --ignorar-apos em Heading).")
    return 0


if __name__ == "__main__":
    sys.exit(main())
