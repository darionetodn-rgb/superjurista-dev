#!/usr/bin/env python3
"""verificar_citacoes.py - Gate determinístico de citações verbatim (Iron Law nº 1).

Extrai cada trecho ENTRE ASPAS do documento e confere, por correspondência
normalizada (acentos/caixa/espaços), se é cópia exata de: (a) um trecho_verbatim
do arquivo de fontes ($NUMERO-fontes.json — cadeia de custódia dos MCPs) ou
(b) dos autos (processo.txt / $NUMERO.txt). Autenticidade vira exit code; a
PERTINÊNCIA (a fonte sustenta a proposição?) continua sendo do verificador-fontes.

Trechos com corte interno ("(...)", "[...]", "…") são divididos: cada fragmento
com >= LIMIAR chars normalizados deve constar do corpus. Aspas curtas
(expressões idiomáticas) são ignoradas.

Uso:
  python scripts/verificar_citacoes.py <workspace> [--id N] [--doc SUFIXO_OU_CAMINHO]
      [--limiar 60] [--ignorar-apos MARCADOR]
      -> exit 0: todas as citações conferem; 1: alguma não confere; 2: erro de uso

LIMIAR default 60 (calibrado contra fundamentações reais — ver Tarefa 3 do plano
2026-07-11). --doc aceita sufixo ("-sentenca.md") ou caminho completo de minuta,
em .docx (Word, lido por python-docx — inclusive texto dentro de TABELA) ou em
texto puro (.md/.txt). Ver ler_documento().
ATENÇÃO: sufixo começa com hífen — usar a forma --doc=-sentenca.md (com "=");
argparse rejeita "--doc -sentenca.md" com valor separado por espaço.

--ignorar-apos MARCADOR trunca o documento na primeira linha de HEADING
Markdown (linha que começa com '#') cujo conteúdo (normalizado de acento/caixa)
contém o marcador: citações a partir dessa linha ficam FORA do regime. A
truncagem é por LINHA inteira — determinística e imune aos deslocamentos de
posição que a normalização causaria num corte por índice — e SÓ em heading:
menção ao marcador em texto corrido (ex.: num sumário executivo antes do
corpo) não trunca. Motivação (calibração T3, 11/07/2026): a minuta-robustecida
AUTO-CITA trechos da própria minuta na seção "## Log de Alterações" (26 falsos
positivos num documento real); no re-gate final do pipeline de revisão usa-se
--ignorar-apos "log de alterações". Um [AVISO] no output registra o marcador e
a linha do corte (ou a ausência do marcador).

Registro de calibração (2026-07-11, corpus = só autos, sem fontes.json):
  9 documentos reais de data/sentenca — fundamentações 0005144-15, 0053258-19,
  0808371-14, 0005006-48, 0076335-57, 0006231-56, 0802594-19; minuta-robustecida
  0802594-19 (via --doc); artefato solto 0808810-30 (via --id). 31 citações no
  regime, 29 [ERRO]: VERDADEIRO=2 (Súmulas 647 e 297 do STJ citadas de memória,
  72–144 chars norm.), FALSO-LEGISLACAO=0 (transcrições de lei observadas medem
  42–57 chars norm. e já ficam fora do regime), FALSO-AUTOS=0 (citações do edital
  casaram com processo.txt real de OCR), FALSO-OUTRO=27 (26 são auto-citações do
  log de alterações da minuta-robustecida — documento fora do alvo do gate — e 1
  é citação retórico-hipotética de 62 chars). Conclusão: LIMIAR mantido em 60
  (separa empiricamente legislação curta <60 de teses/súmulas >=72); nenhuma flag
  adicionada; nenhum ajuste de matcher exigido pelos dados.
"""
import argparse
import json
import os
import re
import sys
import unicodedata

LIMIAR_DEFAULT = 60
RE_CNJ = re.compile(r"\d{7}-\d{2}\.\d{4}\.\d\.\d{2}\.\d{4}")
# pares de aspas: retas duplas e tipográficas; DOTALL para citações multilinha
RE_ASPAS = re.compile(r'"([^"]+)"|“([^”]+)”', re.DOTALL)
RE_CORTE = re.compile(r"\(\s*\.\.\.\s*\)|\[\s*\.\.\.\s*\]|\(\s*…\s*\)|\[\s*…\s*\]|…|\.\.\.")


def _norm(s):
    s = unicodedata.normalize("NFD", s)
    s = "".join(c for c in s if unicodedata.category(c) != "Mn").casefold()
    return re.sub(r"\s+", " ", s).strip()


def carregar_corpus(workspace, ident):
    partes, avisos = [], []
    caminho_fontes = os.path.join(workspace, ident + "-fontes.json")
    if os.path.exists(caminho_fontes):
        with open(caminho_fontes, encoding="utf-8") as f:
            dados = json.load(f)
        for fonte in dados.get("fontes", []):
            partes.append(_norm(fonte.get("trecho_verbatim") or ""))
    else:
        avisos.append("sem " + os.path.basename(caminho_fontes) + " (corpus = só autos)")
    for nome in ("processo.txt", ident + ".txt"):
        p = os.path.join(workspace, nome)
        if os.path.exists(p):
            partes.append(_norm(open(p, encoding="utf-8", errors="replace").read()))
    if len(partes) == 0:
        avisos.append("corpus VAZIO (nem fontes.json nem autos)")
    return "  ".join(partes), avisos


def extrair_citacoes(texto):
    for m in RE_ASPAS.finditer(texto):
        yield (m.group(1) or m.group(2)), m.start()


def truncar_apos_marcador(texto, marcador):
    """Corta o texto na primeira linha de HEADING Markdown (começa com '#')
    cujo conteúdo normalizado contém o marcador normalizado. Retorna
    (texto_truncado, nº da linha do corte | None). Só headings contam: menção
    ao marcador em texto corrido (ex.: num sumário executivo antes do corpo)
    NÃO trunca — truncar ali faria o gate pular o corpo silenciosamente."""
    alvo = _norm(marcador)
    if alvo:
        linhas = texto.splitlines(keepends=True)
        for i, linha in enumerate(linhas):
            if linha.lstrip().startswith("#") and alvo in _norm(linha):
                return "".join(linhas[:i]), i + 1
    return texto, None


def verificar(doc_texto, corpus, limiar):
    """Lista de (trecho_original, fragmento_reprovado|None). Vazia = tudo OK."""
    problemas, conferidas = [], 0
    for bruto, _pos in extrair_citacoes(doc_texto):
        fragmentos = [f for f in RE_CORTE.split(bruto)]
        relevantes = [f for f in fragmentos if len(_norm(f)) >= limiar]
        if not relevantes:
            continue  # aspas curtas: fora do regime
        conferidas += 1
        for frag in relevantes:
            if _norm(frag) not in corpus:
                problemas.append((bruto, frag))
                break
    return problemas, conferidas


def ler_documento(caminho):
    """Lê o documento alvo. Aceita .docx (Word) e texto puro (.md/.txt).

    A minuta do escritório é sempre .docx; até 13/08/2026 este script abria
    qualquer caminho com open(encoding="utf-8") e morria em UnicodeDecodeError,
    e a conversão manual se repetia a cada corrida do /pipeline-revisao-minuta.

    Duas armadilhas conhecidas, tratadas aqui:

    1. `documento.paragraphs` do python-docx NÃO enxerga texto dentro de tabela.
       O corpo é percorrido pelo XML, na ordem real do documento, para que
       citação dentro de tabela entre no regime do gate.
    2. `PackageNotFoundError` é a MESMA exceção para "não existe", "está aberto
       no Word" e "não é um zip/docx real" — o traceback cru manda caçar a causa
       errada. A mensagem abaixo nomeia as três.

    O .docx sai como texto tipo Markdown: parágrafo de estilo "Heading N" vira
    linha iniciada por '#', para que --ignorar-apos (que corta em HEADING) valha
    igual nos dois formatos.
    """
    if not caminho.lower().endswith(".docx"):
        return open(caminho, encoding="utf-8").read()

    try:
        import docx
        from docx.opc.exceptions import PackageNotFoundError
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError:
        print("[ERRO] leitura de .docx exige python-docx: py -3 -m pip install python-docx")
        sys.exit(2)

    try:
        documento = docx.Document(caminho)
    except PackageNotFoundError:
        print(f"[ERRO] python-docx nao abriu: {caminho}\n"
              "       A mesma excecao cobre TRES causas distintas — confira as tres:\n"
              "       (a) o caminho nao existe;\n"
              "       (b) o arquivo esta ABERTO no Word (feche e repita);\n"
              "       (c) nao e um .docx real (arquivo renomeado).")
        sys.exit(2)

    linhas = []
    for filho in documento.element.body.iterchildren():
        if filho.tag.endswith("}p"):
            paragrafo = Paragraph(filho, documento)
            texto = paragrafo.text
            estilo = paragrafo.style.name if paragrafo.style is not None else ""
            if texto.strip() and (estilo or "").lower().startswith("heading"):
                linhas.append("# " + texto)
            else:
                linhas.append(texto)
        elif filho.tag.endswith("}tbl"):
            for linha_tabela in Table(filho, documento).rows:
                for celula in linha_tabela.cells:
                    for p in celula.paragraphs:
                        if p.text.strip():
                            linhas.append(p.text)
    return "\n".join(linhas)


def main():
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass
    ap = argparse.ArgumentParser(description="Gate de citações verbatim (v3.0).")
    ap.add_argument("workspace")
    ap.add_argument("--id", dest="ident", help="prefixo (inferido do nome da pasta)")
    ap.add_argument("--doc", default="-fundamentacao.md",
                    help="sufixo do documento (default -fundamentacao.md) ou caminho")
    ap.add_argument("--limiar", type=int, default=LIMIAR_DEFAULT)
    ap.add_argument("--ignorar-apos", dest="ignorar_apos", metavar="MARCADOR",
                    help="trunca o documento na primeira linha de HEADING ('#'...) que "
                         "contém o marcador (normalizado de acento/caixa); citações após "
                         "essa linha ficam fora do regime (ex.: 'log de alterações')")
    a = ap.parse_args()

    if not os.path.isdir(a.workspace):
        print(f"[ERRO] workspace inexistente: {a.workspace}")
        sys.exit(2)
    base = os.path.basename(os.path.abspath(a.workspace))
    m = RE_CNJ.search(base)
    ident = a.ident or (m.group(0) if m else base)

    doc = a.doc if os.path.sep in a.doc or os.path.exists(a.doc) \
        else os.path.join(a.workspace, ident + a.doc)
    if not os.path.exists(doc):
        print(f"[ERRO] documento inexistente: {doc}")
        sys.exit(2)

    corpus, avisos = carregar_corpus(a.workspace, ident)
    texto = ler_documento(doc)
    if a.ignorar_apos:
        texto, linha_corte = truncar_apos_marcador(texto, a.ignorar_apos)
        if linha_corte:
            avisos.append(f"documento truncado no marcador {a.ignorar_apos!r} (linha {linha_corte})")
        else:
            avisos.append(f"marcador {a.ignorar_apos!r} não encontrado — documento inteiro examinado")
    problemas, conferidas = verificar(texto, corpus, a.limiar)

    print(f"[INICIO] {os.path.basename(doc)} -> {conferidas} citação(ões) no regime verbatim")
    for av in avisos:
        print(f"[AVISO] {av}")
    for bruto, frag in problemas:
        resumo = re.sub(r"\s+", " ", frag).strip()[:120]
        print(f"[ERRO] citação sem lastro no corpus: \"{resumo}...\"")
    print(f"[FIM] {conferidas - len(problemas)}/{conferidas} conferidas")
    sys.exit(1 if problemas else 0)


if __name__ == "__main__":
    main()
